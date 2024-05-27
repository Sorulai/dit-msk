import os
from django.conf import settings
from docx import Document
from datetime import datetime, timedelta
from .models import Booking


def generate_report(start_date, end_date, room_id=None):
    bookings = Booking.objects.filter(start_time__gte=start_date, end_time__lte=end_date)
    if room_id:
        bookings = bookings.filter(room_id=room_id)

    document = Document()
    document.add_heading('Booking Report', 0)
    for booking in bookings:
        document.add_heading(f'Booking by {booking.user.username}', level=1)
        document.add_paragraph(f'Room: {booking.room.name}')
        document.add_paragraph(f'Start Time: {booking.start_time}')
        document.add_paragraph(f'End Time: {booking.end_time}')
        document.add_paragraph(f'Purpose: {booking.purpose}')

    report_name = f'report_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx'
    report_path = os.path.join(settings.MEDIA_ROOT, report_name)
    document.save(report_path)

    return report_name


def processing_dates(room):
    current_date = datetime.now().date()
    start_time = datetime.combine(current_date, datetime.strptime('07:00', '%H:%M').time())
    end_time = datetime.combine(current_date, datetime.strptime('22:00', '%H:%M').time())

    time_slots = []
    current_slot = start_time
    while current_slot <= end_time:
        time_slots.append(current_slot)
        current_slot += timedelta(hours=1)

    bookings = Booking.objects.filter(room=room, start_time__date=current_date, start_time__gte=start_time,
                                      end_time__lte=end_time)
    booked_times = set()
    for booking in bookings:
        current_slot = booking.start_time
        while current_slot < booking.end_time:
            booked_times.add(current_slot)
            current_slot += timedelta(hours=1)

    results = []
    for slot in time_slots:
        result = {
            'start_time': slot,
            'end_time': slot + timedelta(hours=1),
            'booked': slot in booked_times
        }
        results.append(result)

    return results
